import tkinter as tk
from tkinter import filedialog, ttk, messagebox
from docx import Document
import pdfplumber
from deep_translator import GoogleTranslator
from langdetect import detect
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import threading
import os
import time

# OCR
import pytesseract
from PIL import Image

# ---------------- TESSERACT PATH SETUP ----------------
pytesseract.pytesseract.tesseract_cmd = r"E:\SEM-6\AI\tesseract.exe"

# ---------------- CONFIGURATION & COLORS ----------------
COLORS = {
    # Deep dark cinematic palette
    "bg_deep":      "#080c14",       # Almost-black navy
    "bg_panel":     "#0d1525",       # Card / panel background
    "bg_input":     "#111b2e",       # Text-area background
    "border":       "#1e3a5f",       # Subtle blue border
    "border_glow":  "#2563eb",       # Active / focus glow
    "primary":      "#2563eb",       # Electric blue
    "primary_dim":  "#1d4ed8",
    "accent":       "#06d6a0",       # Cyan-green accent
    "accent_dim":   "#059669",
    "danger":       "#ef4444",
    "text_bright":  "#e2f0ff",       # Near-white text
    "text_mid":     "#7ea8d4",       # Muted blue-grey text
    "text_dim":     "#3d6080",       # Dimmer label text
    "gold":         "#f59e0b",       # Detected-language chip
    "white":        "#ffffff",
}

FONT_TITLE   = ("Segoe UI", 22, "bold")
FONT_HEADING = ("Segoe UI", 10, "bold")
FONT_BODY    = ("Segoe UI", 11)
FONT_SMALL   = ("Segoe UI", 9)
FONT_MONO    = ("Consolas", 10)

languages = {
    "English":"en","Gujarati":"gu","Hindi":"hi","French":"fr","Spanish":"es",
    "German":"de","Arabic":"ar","Chinese":"zh-cn","Japanese":"ja",
    "Russian":"ru","Italian":"it","Portuguese":"pt","Korean":"ko",
    "Dutch":"nl","Turkish":"tr"
}

# ─────────────────────────────────────────────
#  CORE LOGIC  (unchanged — only UI wrapping)
# ─────────────────────────────────────────────

def update_word_counts(event=None):
    input_text  = input_box.get("1.0", tk.END).strip()
    input_words = len(input_text.split()) if input_text else 0
    input_count_label.config(text=f"⬡  {input_words} words")

    output_text  = output_box.get("1.0", tk.END).strip()
    output_words = len(output_text.split()) if output_text else 0
    output_count_label.config(text=f"⬡  {output_words} words")

def extract_text_from_file(file_path):
    text = ""
    try:
        if file_path.endswith(".docx"):
            doc  = Document(file_path)
            text = "\n".join([p.text for p in doc.paragraphs])
        elif file_path.endswith(".txt"):
            with open(file_path, "r", encoding="utf8") as f:
                text = f.read()
        elif file_path.endswith(".pdf"):
            with pdfplumber.open(file_path) as pdf:
                text = "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])
    except Exception as e:
        print(f"Error extracting text: {e}")
    return text

def ocr_process():
    file = filedialog.askopenfilename(filetypes=[("Image Files", "*.png *.jpg *.jpeg")])
    if not file: return
    def run_ocr():
        try:
            set_status("Processing OCR…", COLORS["primary"])
            img  = Image.open(file)
            text = pytesseract.image_to_string(img)
            input_box.delete("1.0", tk.END)
            input_box.insert(tk.END, text)
            set_status("OCR Complete", COLORS["accent"])
            update_word_counts()
        except Exception:
            messagebox.showerror("OCR Error", "Tesseract not found.")
            set_status("OCR Failed", COLORS["danger"])
    threading.Thread(target=run_ocr, daemon=True).start()

def translate_logic():
    original = input_box.get("1.0", tk.END).strip()
    if not original:
        messagebox.showwarning("Empty", "Please enter or upload some text first.")
        return
    target_lang = languages.get(language_var.get(), "en")
    progress["value"] = 0
    set_status("Translating…", COLORS["primary"])
    def run():
        try:
            translator  = GoogleTranslator(source='auto', target=target_lang)
            chunk_size  = 3000
            chunks      = [original[i:i+chunk_size] for i in range(0, len(original), chunk_size)]
            full_text   = ""
            total       = len(chunks)
            for i, chunk in enumerate(chunks):
                if i == 0:
                    try:
                        det = detect(chunk)
                        detected_label.config(text=f"  {det.upper()}  ")
                    except:
                        detected_label.config(text="  ??  ")
                full_text       += translator.translate(chunk) + " "
                progress["value"] = int(((i+1)/total)*100)
                root.update_idletasks()
                time.sleep(0.5)
            output_box.delete("1.0", tk.END)
            output_box.insert(tk.END, full_text.strip())
            set_status("Translation Complete", COLORS["accent"])
            update_word_counts()
        except Exception as e:
            messagebox.showerror("Error", f"Translation failed: {e}")
            set_status("Error", COLORS["danger"])
    threading.Thread(target=run, daemon=True).start()

def upload_action():
    files = filedialog.askopenfilenames(filetypes=[("Documents", "*.docx *.txt *.pdf")])
    if not files: return
    set_status(f"Loading {len(files)} file(s)…", COLORS["primary"])
    combined = ""
    for idx, f in enumerate(files):
        name    = os.path.basename(f)
        content = extract_text_from_file(f)
        combined += f"FILE [{idx+1}]: {name}\n" + "="*40 + "\n" + content + "\n\n" + "-"*40 + "\n\n"
    input_box.delete("1.0", tk.END)
    input_box.insert(tk.END, combined.strip())
    set_status(f"{len(files)} File(s) Loaded", COLORS["accent"])
    update_word_counts()

# ── export ──
def save_as_docx():
    content = output_box.get("1.0", tk.END).strip()
    if not content: return
    path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document","*.docx")])
    if path:
        try:
            doc = Document(); doc.add_paragraph(content); doc.save(path)
            messagebox.showinfo("Saved", "Exported to Word!")
        except Exception as e:
            messagebox.showerror("Error", str(e))

def save_as_txt():
    content = output_box.get("1.0", tk.END).strip()
    if not content: return
    path = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[("Text File","*.txt")])
    if path:
        with open(path, "w", encoding="utf-8") as f: f.write(content)
        messagebox.showinfo("Saved", "Exported to Text (UTF-8)!")

def save_as_pdf():
    content = output_box.get("1.0", tk.END).strip()
    if not content: return
    path = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF File","*.pdf")])
    if path:
        try:
            c = canvas.Canvas(path, pagesize=letter)
            width, height  = letter
            guj_font       = "Lohit-Gujarati.ttf"
            hin_font       = "Lohit-Devanagari.ttf"
            current_font   = 'Helvetica'
            target_lang_full = language_var.get()
            if target_lang_full == "Gujarati" and os.path.exists(guj_font):
                pdfmetrics.registerFont(TTFont('GujaratiFont', guj_font))
                current_font = 'GujaratiFont'
            elif target_lang_full == "Hindi" and os.path.exists(hin_font):
                pdfmetrics.registerFont(TTFont('HindiFont', hin_font))
                current_font = 'HindiFont'
            text_obj = c.beginText(50, height-50)
            text_obj.setFont(current_font, 11)
            for line in content.split('\n'):
                words       = line.split(' ')
                current_line = ""
                for word in words:
                    if len(current_line+word) < 75:
                        current_line += word+" "
                    else:
                        text_obj.textLine(current_line); current_line = word+" "
                text_obj.textLine(current_line)
                if text_obj.getY() < 50:
                    c.drawText(text_obj); c.showPage()
                    text_obj = c.beginText(50, height-50)
                    text_obj.setFont(current_font, 11)
            c.drawText(text_obj); c.save()
            messagebox.showinfo("Saved", f"PDF Saved using {current_font}!")
        except Exception as e:
            messagebox.showerror("PDF Error", f"An error occurred: {e}")

# ─────────────────────────────────────────────
#  UI HELPERS
# ─────────────────────────────────────────────

def set_status(msg, color=None):
    color = color or COLORS["text_mid"]
    status_label.config(text=msg, fg=color)

def make_icon_button(parent, icon, label, cmd, bg, hover_bg=None, fg=None):
    """Flat icon+label button with hover effect."""
    hover_bg = hover_bg or bg
    fg       = fg       or COLORS["white"]
    btn = tk.Button(
        parent,
        text=f"{icon}  {label}",
        command=cmd,
        bg=bg, fg=fg,
        activebackground=hover_bg,
        activeforeground=fg,
        font=("Segoe UI", 9, "bold"),
        relief="flat",
        bd=0,
        padx=16, pady=8,
        cursor="hand2",
    )
    btn.bind("<Enter>", lambda e: btn.config(bg=hover_bg))
    btn.bind("<Leave>", lambda e: btn.config(bg=bg))
    return btn

def separator(parent, orient="h", pad=0):
    color = COLORS["border"]
    if orient == "h":
        tk.Frame(parent, bg=color, height=1).pack(fill="x", pady=pad)
    else:
        tk.Frame(parent, bg=color, width=1).pack(fill="y", padx=pad)

# ─────────────────────────────────────────────
#  ROOT WINDOW
# ─────────────────────────────────────────────

root = tk.Tk()
root.title("AI Document Translator Pro")
root.geometry("1200x860")
root.minsize(900, 700)
root.configure(bg=COLORS["bg_deep"])

# ── ttk style for combobox & progressbar ──
style = ttk.Style()
style.theme_use("clam")

style.configure(
    "Dark.TCombobox",
    fieldbackground=COLORS["bg_input"],
    background=COLORS["bg_input"],
    foreground=COLORS["text_bright"],
    arrowcolor=COLORS["primary"],
    bordercolor=COLORS["border"],
    lightcolor=COLORS["border"],
    darkcolor=COLORS["border"],
    selectbackground=COLORS["primary"],
    selectforeground=COLORS["white"],
    padding=(8, 6),
)
style.map("Dark.TCombobox",
    fieldbackground=[("readonly", COLORS["bg_input"])],
    foreground=[("readonly", COLORS["text_bright"])],
)

style.configure(
    "Neon.Horizontal.TProgressbar",
    troughcolor=COLORS["bg_panel"],
    background=COLORS["primary"],
    bordercolor=COLORS["bg_panel"],
    lightcolor=COLORS["primary"],
    darkcolor=COLORS["primary_dim"],
    thickness=4,
)

# ─────────────────────────────────────────────
#  HEADER BAR
# ─────────────────────────────────────────────

header = tk.Frame(root, bg=COLORS["bg_panel"], height=70)
header.pack(fill="x")
header.pack_propagate(False)

# Left: decorative accent bar + title
accent_bar = tk.Frame(header, bg=COLORS["primary"], width=4)
accent_bar.pack(side="left", fill="y")

title_frame = tk.Frame(header, bg=COLORS["bg_panel"])
title_frame.pack(side="left", padx=20)

tk.Label(
    title_frame,
    text="MULTILANGUAGE  TRANSLATOR",
    font=("Segoe UI", 18, "bold"),
    bg=COLORS["bg_panel"],
    fg=COLORS["text_bright"],
).pack(anchor="w")

tk.Label(
    title_frame,
    text="by  Om Bhatt  ·  AI Document Pro",
    font=("Segoe UI", 9),
    bg=COLORS["bg_panel"],
    fg=COLORS["text_dim"],
).pack(anchor="w")

# Right: detected language chip
chip_outer = tk.Frame(header, bg=COLORS["bg_panel"])
chip_outer.pack(side="right", padx=24)

tk.Label(chip_outer, text="DETECTED", font=("Segoe UI", 7, "bold"),
         bg=COLORS["bg_panel"], fg=COLORS["text_dim"]).pack(anchor="e")

detected_label = tk.Label(
    chip_outer,
    text="  NONE  ",
    font=("Segoe UI", 11, "bold"),
    bg=COLORS["gold"],
    fg=COLORS["bg_deep"],
    padx=6, pady=2,
)
detected_label.pack()

separator(root, pad=0)

# ─────────────────────────────────────────────
#  TOOLBAR
# ─────────────────────────────────────────────

toolbar = tk.Frame(root, bg=COLORS["bg_panel"], pady=12)
toolbar.pack(fill="x")

tk.Label(
    toolbar,
    text="TARGET LANGUAGE",
    font=("Segoe UI", 7, "bold"),
    bg=COLORS["bg_panel"],
    fg=COLORS["text_dim"],
).pack(side="left", padx=(20, 4))

language_var = tk.StringVar(value="Gujarati")
lang_cb = ttk.Combobox(
    toolbar,
    textvariable=language_var,
    values=list(languages.keys()),
    state="readonly",
    width=14,
    style="Dark.TCombobox",
)
lang_cb.pack(side="left", padx=(0, 18))

# Divider
tk.Frame(toolbar, bg=COLORS["border"], width=1, height=30).pack(side="left", padx=6)

make_icon_button(toolbar, "📁", "Upload Files", upload_action,
                 bg="#1e3a5f", hover_bg="#2563eb").pack(side="left", padx=6)

make_icon_button(toolbar, "📷", "Image OCR", ocr_process,
                 bg="#1e3a5f", hover_bg="#7c3aed").pack(side="left", padx=6)

tk.Frame(toolbar, bg=COLORS["border"], width=1, height=30).pack(side="left", padx=6)

make_icon_button(toolbar, "⚡", "Translate Now", translate_logic,
                 bg=COLORS["accent_dim"], hover_bg=COLORS["accent"],
                 fg=COLORS["bg_deep"]).pack(side="left", padx=6)

separator(root, pad=0)

# ─────────────────────────────────────────────
#  PROGRESS + STATUS
# ─────────────────────────────────────────────

prog_frame = tk.Frame(root, bg=COLORS["bg_deep"], pady=6)
prog_frame.pack(fill="x", padx=24)

progress = ttk.Progressbar(
    prog_frame,
    length=900,
    mode="determinate",
    style="Neon.Horizontal.TProgressbar",
)
progress.pack(fill="x")

status_label = tk.Label(
    root,
    text="Ready",
    font=("Segoe UI", 9),
    bg=COLORS["bg_deep"],
    fg=COLORS["text_dim"],
    anchor="w",
)
status_label.pack(fill="x", padx=28)

# ─────────────────────────────────────────────
#  MAIN TEXT AREA
# ─────────────────────────────────────────────

content_frame = tk.Frame(root, bg=COLORS["bg_deep"])
content_frame.pack(expand=True, fill="both", padx=20, pady=(8, 0))

def create_editor_panel(parent, title, readonly=False):
    """Returns (outer_frame, text_widget, word_count_label)"""
    outer = tk.Frame(parent, bg=COLORS["bg_panel"], bd=0,
                     highlightbackground=COLORS["border"], highlightthickness=1)

    # Panel header
    ph = tk.Frame(outer, bg=COLORS["bg_panel"], pady=8)
    ph.pack(fill="x", padx=12)

    tk.Label(ph, text=title, font=("Segoe UI", 9, "bold"),
             bg=COLORS["bg_panel"], fg=COLORS["primary"]).pack(side="left")

    wc = tk.Label(ph, text="⬡  0 words", font=("Segoe UI", 8),
                  bg=COLORS["bg_panel"], fg=COLORS["text_dim"])
    wc.pack(side="right")

    tk.Frame(outer, bg=COLORS["border"], height=1).pack(fill="x")

    # Text + scrollbar
    txt_frame = tk.Frame(outer, bg=COLORS["bg_input"])
    txt_frame.pack(expand=True, fill="both")

    sb = tk.Scrollbar(txt_frame, bg=COLORS["bg_panel"],
                      troughcolor=COLORS["bg_input"], relief="flat", bd=0)
    sb.pack(side="right", fill="y")

    txt = tk.Text(
        txt_frame,
        font=("Segoe UI", 11),
        bg=COLORS["bg_input"],
        fg=COLORS["text_bright"],
        insertbackground=COLORS["accent"],
        selectbackground=COLORS["primary"],
        selectforeground=COLORS["white"],
        relief="flat",
        bd=0,
        padx=14, pady=12,
        wrap="word",
        yscrollcommand=sb.set,
        cursor="xterm",
    )
    if readonly:
        txt.configure(bg="#0f1f38")
    txt.pack(expand=True, fill="both")
    sb.configure(command=txt.yview)
    txt.bind("<KeyRelease>", update_word_counts)

    return outer, txt, wc

in_panel,  input_box,  input_count_label  = create_editor_panel(content_frame, "◈  SOURCE TEXT")
out_panel, output_box, output_count_label = create_editor_panel(content_frame, "◈  TRANSLATED TEXT", readonly=True)

in_panel.pack(side="left",  expand=True, fill="both", padx=(0, 8))
out_panel.pack(side="right", expand=True, fill="both", padx=(8, 0))

# ─────────────────────────────────────────────
#  FOOTER  (export buttons)
# ─────────────────────────────────────────────

separator(root, pad=0)

footer = tk.Frame(root, bg=COLORS["bg_panel"], pady=14)
footer.pack(fill="x", side="bottom")

tk.Label(footer, text="EXPORT", font=("Segoe UI", 7, "bold"),
         bg=COLORS["bg_panel"], fg=COLORS["text_dim"]).pack(side="left", padx=(24, 8))

make_icon_button(footer, "📄", "Word .docx", save_as_docx,
                 bg="#1e3a5f", hover_bg="#2b579a").pack(side="left", padx=5)

make_icon_button(footer, "📕", "PDF", save_as_pdf,
                 bg="#3b0a0a", hover_bg="#b91c1c").pack(side="left", padx=5)

make_icon_button(footer, "🗒️", "Plain Text", save_as_txt,
                 bg="#1a2e1a", hover_bg="#166534").pack(side="left", padx=5)

tk.Label(footer, text="Developed by Om Bhatt  ·  AI Document Translator Pro",
         font=("Segoe UI", 8, "italic"),
         bg=COLORS["bg_panel"],
         fg=COLORS["text_dim"]).pack(side="right", padx=24)

# ─────────────────────────────────────────────
root.mainloop()